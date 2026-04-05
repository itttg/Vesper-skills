#!/usr/bin/env python3
"""Fill OCB repayment schedule template."""
import argparse
from datetime import date
from docx import Document


def fmt_amount(value: int) -> str:
    return f"{value:,.0f}".replace(",", ".")


def replace_all(doc: Document, old: str, new: str) -> None:
    for para in doc.paragraphs:
        if old in para.text:
            para.text = para.text.replace(old, new)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if old in para.text:
                        para.text = para.text.replace(old, new)


def build_installments(total_amount: int, periods: int = 12):
    base = total_amount // periods
    remainder = total_amount - (base * periods)
    amounts = [base] * periods
    amounts[-1] += remainder
    return amounts


def fill_schedule(template_path: str, output_path: str, total_amount: int, kunn_date: str, year: int | None = None):
    doc = Document(template_path)
    today = date.today()
    year = year or today.year

    replace_all(doc, '/2025/KUNN-OCB-DN', f'/{year}/KUNN-OCB-DN')
    replace_all(doc, '/2026/KUNN-OCB-DN', f'/{year}/KUNN-OCB-DN')
    replace_all(doc, 'ngày …../…../………', f'ngày {kunn_date}')
    replace_all(doc, 'Tổng số tiền phải trả: ……', f'Tổng số tiền phải trả: {fmt_amount(total_amount)} đồng')

    installments = build_installments(total_amount, 12)
    table = doc.tables[0]
    for idx, amount in enumerate(installments, start=1):
        table.rows[idx].cells[2].text = fmt_amount(amount)

    doc.save(output_path)
    return {
        'total_amount': fmt_amount(total_amount),
        'installments': [fmt_amount(x) for x in installments],
        'kunn_date': kunn_date,
    }


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument('template')
    parser.add_argument('output')
    parser.add_argument('--tong-so-tien', type=int, required=True)
    parser.add_argument('--ngay-kunn', required=True)
    parser.add_argument('--nam', type=int, default=None)
    args = parser.parse_args()

    info = fill_schedule(args.template, args.output, args.tong_so_tien, args.ngay_kunn, args.nam)
    print('✅ Filled repayment schedule:')
    print(f"   Total amount: {info['total_amount']} đồng")
    print(f"   Installment count: {len(info['installments'])}")
    print(f"   Saved to: {args.output}")


if __name__ == '__main__':
    main()
