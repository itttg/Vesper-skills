#!/usr/bin/env python3
"""Fill OCB supporting documents list from payment package PDFs."""
import argparse
import os
import re
import subprocess
from datetime import date
from pathlib import Path
from docx import Document


def fmt_amount(value: int | None) -> str:
    if value is None:
        return ''
    return f"{value:,.0f}".replace(',', '.')


def replace_all(doc: Document, old: str, new: str) -> None:
    for para in doc.paragraphs:
        if old in para.text:
            para.text = para.text.replace(old, new)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if old in para.text:
                        para.text = para.text.replace(old, new)


def extract_pdf_text(pdf_path: Path) -> str:
    try:
        result = subprocess.run(
            ['pdftotext', '-layout', str(pdf_path), '-'],
            capture_output=True,
            text=True,
            encoding='utf-8',
            errors='ignore',
            check=False,
        )
        return result.stdout or ''
    except Exception:
        return ''


def normalize_spaces(text: str) -> str:
    return re.sub(r'\s+', ' ', text).strip()


def parse_payment_request_text(text: str, filename: str, total_amount: int) -> dict:
    compact = normalize_spaces(text)
    vendor = ''
    vendor_match = re.search(r'(Công ty TNHH[^\n\r]+|Cong ty TNHH[^\n\r]+)', text, re.IGNORECASE)
    if vendor_match:
        vendor = normalize_spaces(vendor_match.group(1))
    elif 'TK STUDIO' in compact.upper():
        vendor = 'Công ty TNHH TK STUDIO'

    contract_no = ''
    contract_match = re.search(r'([A-Z0-9_\-]*TK STUDIO[A-Z0-9_\-]*)', compact, re.IGNORECASE)
    if contract_match:
        contract_no = contract_match.group(1).strip()

    contract_date = ''
    all_dates = re.findall(r'(\d{2}/\d{2}/\d{4})', compact)
    if len(all_dates) >= 2:
        contract_date = all_dates[1]
    elif all_dates:
        contract_date = all_dates[0]

    content_line = ''
    for line in text.splitlines():
        if 'IPR' in line.upper() or 'Tender Award' in line:
            content_line = normalize_spaces(line)
            break
    if not content_line:
        content_line = 'Supporting document extracted from payment request package'

    doc_label_parts = []
    if contract_no:
        doc_label_parts.append(f'Hợp đồng số {contract_no}')
    if contract_date:
        doc_label_parts.append(f'ngày {contract_date}')
    if not doc_label_parts:
        doc_label_parts.append(Path(filename).stem)

    return {
        'document_name': ' '.join(doc_label_parts),
        'document_value': total_amount,
        'paid_by_ocb': 0,
        'paid_by_other': 0,
        'remaining_finance': total_amount,
        'current_disbursement': total_amount,
        'payment_content': content_line,
        'beneficiary': vendor,
        'bank_account': '',
        'bank_name': '',
    }


def build_supporting_rows(folder: Path, total_amount: int):
    pdfs = sorted([p for p in folder.glob('*.pdf') if not p.name.startswith('~$')])
    rows = []
    payment_request_added = False
    for pdf in pdfs:
        text = extract_pdf_text(pdf)
        upper = text.upper()
        if 'PAYMENT REQUEST FORM' in upper or 'PHIU  NGH' in upper or 'TK STUDIO' in upper:
            rows.append(parse_payment_request_text(text, pdf.name, total_amount))
            payment_request_added = True
        else:
            rows.append({
                'document_name': Path(pdf.name).stem,
                'document_value': None,
                'paid_by_ocb': None,
                'paid_by_other': None,
                'remaining_finance': None,
                'current_disbursement': None,
                'payment_content': f'Attached supporting file: {pdf.name}',
                'beneficiary': '',
                'bank_account': '',
                'bank_name': '',
            })

    if not payment_request_added:
        rows.insert(0, {
            'document_name': 'Payment package summary',
            'document_value': total_amount,
            'paid_by_ocb': 0,
            'paid_by_other': 0,
            'remaining_finance': total_amount,
            'current_disbursement': total_amount,
            'payment_content': 'Derived from payment package folder',
            'beneficiary': '',
            'bank_account': '',
            'bank_name': '',
        })
    return rows


def set_row_text(row, values):
    for idx, value in enumerate(values):
        row.cells[idx].text = value


def remove_row(table, row):
    table._tbl.remove(row._tr)


def fill_supporting_documents(template_path: str, output_path: str, input_folder: str, total_amount: int, kunn_date: str, year: int | None = None):
    doc = Document(template_path)
    today = date.today()
    year = year or today.year

    replace_all(doc, '/2025/KUNN-OCB-DN', f'/{year}/KUNN-OCB-DN')
    replace_all(doc, '/2026/KUNN-OCB-DN', f'/{year}/KUNN-OCB-DN')
    replace_all(doc, 'ngày ……/…../………', f'ngày {kunn_date}')
    replace_all(doc, 'Đơn vị tính: đồng', 'Đơn vị tính: đồng')

    rows_data = build_supporting_rows(Path(input_folder), total_amount)
    table = doc.tables[0]

    while len(table.rows) > 2:
        remove_row(table, table.rows[2])

    for i, item in enumerate(rows_data, start=1):
        row = table.add_row()
        set_row_text(row, [
            str(i),
            item['document_name'],
            fmt_amount(item['document_value']),
            fmt_amount(item['paid_by_ocb']),
            fmt_amount(item['paid_by_other']),
            fmt_amount(item['remaining_finance']),
            fmt_amount(item['current_disbursement']),
            item['payment_content'],
            item['beneficiary'],
            item['bank_account'],
            item['bank_name'],
        ])

    doc.save(output_path)
    return rows_data


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument('template')
    parser.add_argument('output')
    parser.add_argument('--input-folder', required=True)
    parser.add_argument('--tong-so-tien', type=int, required=True)
    parser.add_argument('--ngay-kunn', required=True)
    parser.add_argument('--nam', type=int, default=None)
    args = parser.parse_args()

    rows = fill_supporting_documents(args.template, args.output, args.input_folder, args.tong_so_tien, args.ngay_kunn, args.nam)
    print('✅ Filled supporting documents list:')
    print(f'   Row count: {len(rows)}')
    print(f'   Saved to: {args.output}')


if __name__ == '__main__':
    main()
